from PyPDF2 import PdfReader
import docx
import pandas as pd
import re

MAX_PDF_PAGES = 20
MAX_EXCEL_ROWS = 1000

def load_document(file):
    file_type = file.name.split(".")[-1].lower()

    try:
        if file_type == "pdf":
            reader = PdfReader(file)
            text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        elif file_type == "docx":
            doc = docx.Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file_type == "xlsx":
            df = pd.read_excel(file)
            text = df.astype(str).apply(lambda x: ' '.join(x), axis=1).str.cat(sep='\n')
        elif file_type == "txt":
            text = file.read().decode("utf-8")
        else:
            return "❌ Unsupported file format."
    except Exception as e:
        return f"❌ Error reading file: {e}"

    if not text.strip():
        return "❌ Document is empty or unreadable."

    return text

def read_pdf(file):
    try:
        reader = PdfReader(file)
        if len(reader.pages) > MAX_PDF_PAGES:
            return f"❌ File too large. Only PDFs with up to {MAX_PDF_PAGES} pages are supported."
        text = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text
        return clean_text(text)
    except Exception as e:
        return f"❌ Failed to read PDF: {e}"

def read_docx(file):
    try:
        doc = docx.Document(file)
        raw_text = "\n".join([para.text for para in doc.paragraphs])
        return clean_text(raw_text)
    except Exception as e:
        return f"❌ Failed to read DOCX: {e}"

def read_excel(file):
    try:
        df = pd.read_excel(file)
        if len(df) > MAX_EXCEL_ROWS:
            return f"❌ Excel file too large. Max {MAX_EXCEL_ROWS} rows allowed."
        return df.to_string(index=False)
    except Exception as e:
        return f"❌ Failed to read Excel: {e}"

def read_txt(file):
    try:
        return file.read().decode("utf-8")
    except Exception as e:
        return f"❌ Failed to read TXT file: {e}"

def clean_text(text: str) -> str:
    text = re.sub(r'\n+', '\n', text)                 # Collapse multiple newlines
    text = re.sub(r'\s{2,}', ' ', text)               # Collapse multiple spaces
    text = re.sub(r'–', '-', text)                    # Normalize dashes
    text = re.sub(r'\u200b', '', text)                # Remove zero-width space
    text = re.sub(r'\n\s*\n', '\n', text)             # Remove empty paragraphs
    return text.strip()
